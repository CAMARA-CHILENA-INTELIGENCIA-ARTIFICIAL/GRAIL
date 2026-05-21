"""
Source-file preprocessing.

Provided by Nirvai (Nirvana). Author: Benjamin González Guerrero.

Turns whatever the user dropped into ``input/`` into the plain-text / markdown
form that :class:`grail.indexing.loader.FileLoader` chunks. The output is
written to ``input/_processed/<name>.md`` so users can inspect or hand-edit it,
and the chunker reads from there.

Supported source types (v0.1):

* **Text-like** (``.txt``, ``.md``, ``.rst``, ``.log``, …) — pass-through.
* **Code** (``.py``, ``.js``, ``.ts``, ``.go``, … — see :mod:`grail.utils.text`) — pass-through.
* **Data** (``.json``, ``.yaml``, ``.csv``, …) — pass-through.
* **PDF** (``.pdf``) — text per page via ``pypdf``. Pages with no extractable
  text are flagged in the log; OCR fallback for image-only pages is on the roadmap.
* **DOCX** (``.docx``) — paragraphs and tables via ``python-docx``. Headings are
  converted to markdown ``#`` headers; tables are flattened to ``|``-separated rows.

The preprocessor is **idempotent** and **cache-aware**: if a processed file
already exists and is newer than the source, the cached version is reused.
Re-run is cheap.

Ported from the legacy ``AgentProcess.py`` PdfExtractor / DocxExtractor with
the S3 / vision-OCR fallback paths removed — those require the proprietary
vision pipeline and aren't needed for the open-source baseline.
"""
from __future__ import annotations

import logging
from abc import ABC, abstractmethod
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

log = logging.getLogger(__name__)


# ----------------------------------------------------------------------- registry


# Extensions that the chunker can read directly (no conversion needed).
DIRECT_READ_EXTENSIONS: frozenset[str] = frozenset({
    # Plain text
    ".txt", ".md", ".markdown", ".rst", ".log", ".srt", ".vtt",
    # Code (any-script support per the user's spec)
    ".py", ".ipynb", ".js", ".ts", ".tsx", ".jsx", ".go", ".rs", ".rb", ".php",
    ".java", ".kt", ".swift", ".c", ".cc", ".cpp", ".h", ".hpp", ".cs", ".scala",
    ".sh", ".bash", ".zsh", ".sql", ".html", ".css", ".scss", ".vue",
    # Structured data
    ".json", ".jsonl", ".yaml", ".yml", ".toml", ".csv", ".tsv", ".xml",
})

# Extensions that require a conversion pass.
PREPROCESS_EXTENSIONS: frozenset[str] = frozenset({".pdf", ".docx", ".doc"})

SUPPORTED_EXTENSIONS: frozenset[str] = DIRECT_READ_EXTENSIONS | PREPROCESS_EXTENSIONS


def needs_preprocessing(path: str | Path) -> bool:
    """Return True if ``path`` requires conversion to text/markdown before chunking."""
    return Path(path).suffix.lower() in PREPROCESS_EXTENSIONS


def is_supported(path: str | Path) -> bool:
    """Return True if GRAIL can ingest ``path`` (with or without preprocessing)."""
    return Path(path).suffix.lower() in SUPPORTED_EXTENSIONS


# ----------------------------------------------------------------------- interface


class Preprocessor(ABC):
    """Convert one source file into plain text or markdown.

    Implementations should:

    * Be deterministic given the same input — so the cache check works.
    * Return a string; never write to disk directly. The caller owns persistence.
    * Raise :class:`PreprocessingError` (or a subclass) on irrecoverable problems.
    """

    @abstractmethod
    def extract(self, source: Path) -> str:
        """Read ``source`` and return its textual content."""

    @property
    def output_extension(self) -> str:
        """File suffix to use when persisting the result. Default ``.md``."""
        return ".md"


class PreprocessingError(RuntimeError):
    """Raised when preprocessing fails for a recoverable reason (no text, malformed file)."""


# ----------------------------------------------------------------------- impls


class TextPreprocessor(Preprocessor):
    """Pass-through for files that are already text-like."""

    @property
    def output_extension(self) -> str:
        return ".txt"

    def extract(self, source: Path) -> str:
        try:
            return source.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return source.read_bytes().decode("utf-8", errors="replace")


class PdfPreprocessor(Preprocessor):
    """Extract text from a PDF via ``pypdf``.

    Each page becomes a ``## Page N`` markdown section. Pages with no extractable
    text are skipped and counted in the log (could be scanned images — OCR is a
    later enhancement).
    """

    def extract(self, source: Path) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as exc:  # pragma: no cover - import guard
            raise ImportError(
                "PDF preprocessing requires `pypdf`. "
                "Install with: pip install pypdf  (already in GRAIL's core deps)."
            ) from exc

        reader = PdfReader(str(source))
        parts: list[str] = []
        empty_pages: list[int] = []
        for i, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if not text or text == "\x0c":
                empty_pages.append(i)
                continue
            parts.append(f"## Page {i}\n\n{text}")

        if not parts:
            raise PreprocessingError(
                f"{source.name} has no extractable text. May be a scanned/image-only PDF — "
                "OCR fallback is on the roadmap."
            )
        if empty_pages:
            log.warning(
                "%s: %d page(s) had no extractable text — skipped: %s",
                source.name,
                len(empty_pages),
                empty_pages if len(empty_pages) <= 10 else f"{empty_pages[:10]}…",
            )
        title = source.stem.replace("_", " ").strip()
        return f"# {title}\n\n" + "\n\n".join(parts)


class DocxPreprocessor(Preprocessor):
    """Extract paragraphs and tables from a DOCX via ``python-docx``.

    Heading styles in the document map to markdown ``#`` levels. Tables are
    flattened to pipe-separated rows. Empty paragraphs are dropped.
    """

    def extract(self, source: Path) -> str:
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover - import guard
            raise ImportError(
                "DOCX preprocessing requires `python-docx`. "
                "Install with: pip install python-docx  (already in GRAIL's core deps)."
            ) from exc

        doc = Document(str(source))
        out: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = ""
            try:
                style_name = (para.style.name if para.style else "") or ""
            except Exception:  # pragma: no cover - defensive
                style_name = ""
            if style_name.startswith("Heading"):
                digits = "".join(ch for ch in style_name if ch.isdigit())
                level = max(1, min(int(digits) if digits else 1, 6))
                out.append(f"{'#' * level} {text}")
            else:
                out.append(text)

        for i, table in enumerate(doc.tables, start=1):
            out.append("")
            out.append(f"### Table {i}")
            for row in table.rows:
                cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
                out.append(" | ".join(cells))

        if not any(line.strip() for line in out):
            raise PreprocessingError(f"{source.name} contains no extractable text.")

        title = source.stem.replace("_", " ").strip()
        return f"# {title}\n\n" + "\n\n".join(out)


# ----------------------------------------------------------------------- factory


def get_preprocessor(extension: str) -> Optional[Preprocessor]:
    """Return the right preprocessor for ``extension``, or ``None`` if unsupported."""
    ext = extension.lower()
    if not ext.startswith("."):
        ext = "." + ext
    if ext in {".pdf"}:
        return PdfPreprocessor()
    if ext in {".docx", ".doc"}:
        return DocxPreprocessor()
    if ext in DIRECT_READ_EXTENSIONS:
        return TextPreprocessor()
    return None


# ----------------------------------------------------------------------- pipeline


@dataclass
class PreprocessResult:
    """Outcome of preprocessing one file."""

    source: Path
    processed: Path
    cached: bool
    error: Optional[str] = None

    @property
    def ok(self) -> bool:
        return self.error is None


def preprocess_file(
    source: Path,
    *,
    output_dir: Path,
    force: bool = False,
) -> PreprocessResult:
    """Convert one file. Persists the result under ``output_dir`` and returns the path.

    The output name is ``<source.stem>.<preprocessor.output_extension>``. Re-runs
    are skipped when the cached file is newer than the source (``force=True``
    bypasses the cache).
    """
    pp = get_preprocessor(source.suffix)
    if pp is None:
        return PreprocessResult(
            source=source,
            processed=source,
            cached=False,
            error=f"Unsupported extension: {source.suffix!r}",
        )

    output_dir.mkdir(parents=True, exist_ok=True)
    dest = output_dir / f"{source.stem}{pp.output_extension}"

    if dest.exists() and not force:
        try:
            if dest.stat().st_mtime >= source.stat().st_mtime:
                return PreprocessResult(source=source, processed=dest, cached=True)
        except OSError:
            pass

    try:
        text = pp.extract(source)
    except PreprocessingError as exc:
        return PreprocessResult(source=source, processed=dest, cached=False, error=str(exc))
    except Exception as exc:  # pragma: no cover - last-resort guard
        return PreprocessResult(
            source=source, processed=dest, cached=False, error=f"{type(exc).__name__}: {exc}"
        )

    dest.write_text(text, encoding="utf-8")
    return PreprocessResult(source=source, processed=dest, cached=False)


def preprocess_directory(
    input_dir: Path,
    *,
    output_subdir: str = "_processed",
    force: bool = False,
) -> list[PreprocessResult]:
    """Preprocess every supported file in ``input_dir``.

    PDFs and DOCX files are extracted into ``input_dir / output_subdir``. Files
    already in a directly-readable format are returned with their original path
    (no copy made).
    """
    output_dir = input_dir / output_subdir
    results: list[PreprocessResult] = []
    for source in sorted(input_dir.iterdir()):
        if not source.is_file():
            continue
        if source.name.startswith(".") or source.name.startswith("_"):
            continue
        ext = source.suffix.lower()
        if ext in DIRECT_READ_EXTENSIONS:
            results.append(PreprocessResult(source=source, processed=source, cached=True))
            continue
        if ext in PREPROCESS_EXTENSIONS:
            results.append(preprocess_file(source, output_dir=output_dir, force=force))
            continue
        log.debug("Skipping %s — unsupported extension", source.name)
    return results
